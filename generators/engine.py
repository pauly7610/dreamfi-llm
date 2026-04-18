"""
Core Document Generation Engine
Uses Anthropic Claude API + templates + DreamFi voice rules.
"""

from __future__ import annotations

import importlib
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any

import anthropic
import markdown as md
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import config as _cfg
from generators.voice import apply_voice_check, VOICE_GUIDELINES

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Registry of available generator templates
# ---------------------------------------------------------------------------

GENERATOR_REGISTRY: dict[str, dict[str, str]] = {
    "technical_prd": {
        "module": "generators.templates.technical_prd",
        "title": "Technical PRD",
        "description": "Product requirements document with system design, API changes, and data model details.",
        "estimated_time": "10-15 min",
    },
    "business_prd": {
        "module": "generators.templates.business_prd",
        "title": "Business PRD",
        "description": "Business-focused PRD with market context, revenue projections, and stakeholder alignment.",
        "estimated_time": "10-15 min",
    },
    "risk_brd": {
        "module": "generators.templates.risk_brd",
        "title": "Third-Party Risk BRD",
        "description": "Risk assessment document for third-party vendor integrations and compliance review.",
        "estimated_time": "15-20 min",
    },
    "sponsor_bank": {
        "module": "generators.templates.sponsor_bank",
        "title": "Sponsor Bank Requirements",
        "description": "Requirements document for sponsor bank integration, settlement, and compliance controls.",
        "estimated_time": "15-20 min",
    },
    "discovery": {
        "module": "generators.templates.discovery",
        "title": "Discovery Document",
        "description": "Research-driven discovery document with hypotheses, evidence gathering, and recommendations.",
        "estimated_time": "10-15 min",
    },
    "epic_builder": {
        "module": "generators.templates.epic_builder",
        "title": "Epic Builder",
        "description": "Generates epic summaries, user stories with acceptance criteria, and dependency maps for Jira.",
        "estimated_time": "8-12 min",
    },
}


def _load_template_module(generator_type: str) -> Any:
    """Dynamically import a template module from the registry."""
    info = GENERATOR_REGISTRY.get(generator_type)
    if not info:
        raise ValueError(f"Unknown generator type: {generator_type}")
    return importlib.import_module(info["module"])


def _build_voice_system_block() -> str:
    """Return a system-prompt fragment encoding DreamFi voice rules."""
    tone_rules = "\n".join(f"- {r}" for r in VOICE_GUIDELINES["tone"]["rules"])
    fmt_rules = "\n".join(f"- {r}" for r in VOICE_GUIDELINES["formatting"]["rules"])
    banned = ", ".join(VOICE_GUIDELINES["banned_phrases"])

    return (
        "## DreamFi Voice Rules\n"
        "Follow these rules in all generated content:\n\n"
        "### Tone\n"
        f"{tone_rules}\n\n"
        "### Formatting\n"
        f"{fmt_rules}\n\n"
        "### Banned Phrases (never use these)\n"
        f"{banned}\n"
    )


# ---------------------------------------------------------------------------
# Generator engine
# ---------------------------------------------------------------------------

class GeneratorEngine:
    """
    Generates DreamFi documents from form data using LLM + templates.

    Usage::

        engine = GeneratorEngine()
        result = engine.generate("technical_prd", form_data)
        docx_bytes = engine.export_to_docx(result["content"], "my_prd.docx")
    """

    def __init__(self):
        self._client = anthropic.Anthropic(api_key=_cfg.anthropic.api_key)
        self._model = _cfg.anthropic.model

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, generator_type: str, form_data: dict[str, Any]) -> dict[str, Any]:
        """
        Generate a document.

        Returns dict with:
            content: str — Markdown document
            voice_check: dict — result of apply_voice_check
            generator_type: str
            title: str
        """
        mod = _load_template_module(generator_type)

        # Build the user prompt from the template
        user_prompt = mod.TEMPLATE.format(**form_data)

        # Build system prompt: voice rules + template-specific instructions
        system_prompt = _build_voice_system_block() + "\n\n" + mod.SYSTEM_PROMPT

        log.info("Generating %s: '%s'", generator_type, form_data.get("title", form_data.get("epic_title", "untitled")))

        response = self._client.messages.create(
            model=self._model,
            max_tokens=8192,
            system=system_prompt,
            messages=[
                {"role": "user", "content": user_prompt},
            ],
        )

        content = response.content[0].text

        voice_check = apply_voice_check(content)

        return {
            "content": content,
            "voice_check": voice_check,
            "generator_type": generator_type,
            "title": form_data.get("title", form_data.get("epic_title", "Generated Document")),
        }

    def export_to_docx(self, content: str, filename: str) -> bytes:
        """
        Convert Markdown content to a .docx file.
        Returns the raw bytes of the docx.
        """
        doc = DocxDocument()

        # Document style setup
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Parse markdown into docx elements
        _markdown_to_docx(doc, content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def publish_to_confluence(
        self,
        content: str,
        title: str,
        space_key: str,
        parent_id: str | None = None,
    ) -> dict[str, Any]:
        """
        Publish Markdown content to Confluence as a new page.
        Converts Markdown to Confluence-compatible HTML first.
        """
        from integrations.confluence_client import ConfluenceClient

        html_body = md.markdown(
            content,
            extensions=["tables", "fenced_code", "toc"],
        )

        # Wrap in Confluence storage-format macro
        confluence_html = (
            '<ac:structured-macro ac:name="html">'
            "<ac:plain-text-body><![CDATA["
            f"{html_body}"
            "]]></ac:plain-text-body>"
            "</ac:structured-macro>"
        )

        client = ConfluenceClient()
        result = client.create_page(
            space_key=space_key,
            title=title,
            body_html=confluence_html,
            parent_id=parent_id,
        )

        log.info("Published '%s' to Confluence space %s", title, space_key)
        return result

    def get_fields(self, generator_type: str) -> list[dict[str, Any]]:
        """Return the FIELDS definition for a generator type."""
        mod = _load_template_module(generator_type)
        return mod.FIELDS

    def get_info(self, generator_type: str) -> dict[str, str]:
        """Return registry metadata for a generator type."""
        return GENERATOR_REGISTRY[generator_type]


# ---------------------------------------------------------------------------
# Markdown-to-docx helper
# ---------------------------------------------------------------------------

def _markdown_to_docx(doc: DocxDocument, content: str) -> None:
    """
    Simple Markdown-to-docx converter.
    Handles headers, bold, bullet lists, tables, and plain paragraphs.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headers
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Table — collect all contiguous | lines
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip()
                # Skip separator rows
                if re.match(r"^\|[\s\-:|]+\|$", row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                table_lines.append(cells)
                i += 1

            if table_lines:
                num_cols = max(len(r) for r in table_lines)
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = "Light Grid Accent 1"
                for row_cells in table_lines:
                    row = table.add_row()
                    for j, cell_text in enumerate(row_cells):
                        if j < num_cols:
                            row.cells[j].text = cell_text
            continue

        # Bullet list
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s", line):
            text = re.sub(r"^\s*\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph, handling **bold** and *italic* markers."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
